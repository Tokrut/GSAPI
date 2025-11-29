"""
Генераторы отчетов для GEO Analyzer Pro
"""

import base64
import pandas as pd
from datetime import datetime
from config import OPTIMAL_VALUES, COLOR_SCHEME
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots
import io

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def safe_get(data, keys, default="Не доступно"):
    """Безопасное получение значения из словаря"""
    try:
        for key in keys:
            if not isinstance(data, dict):
                return default
            data = data.get(key, default)
        return data
    except (KeyError, TypeError, AttributeError, IndexError):
        return default

def generate_full_report(result):
    """Генерация полного текстового отчета с LLM-анализом"""
    
    # Создаем визуализации
    charts_html = generate_charts_html(result)
    
    # Раздел LLM-анализа
    llm_section = generate_llm_analysis_section(result)
    
    report = f"""
GEO ANALYZER PRO - ПОЛНЫЙ ОТЧЕТ АНАЛИЗА САЙТА С LLM-АНАЛИЗОМ
===========================================================
Дата анализа: {safe_get(result, ['basic_info', 'analysis_date'])}
Анализируемый URL: {safe_get(result, ['basic_info', 'url'])}
Финальный URL: {safe_get(result, ['basic_info', 'final_url'], safe_get(result, ['basic_info', 'url']))}
Общая оценка: {safe_get(result, ['score'], 0)}/100
Оценка GEO (LLM): {safe_get(result, ['llm_analysis', 'overall_geo_score'], 'N/A')}/100
Статус: {'Отлично' if safe_get(result, ['score'], 0) >= 80 else 'Хорошо' if safe_get(result, ['score'], 0) >= 60 else 'Требует улучшений' if safe_get(result, ['score'], 0) >= 40 else 'Критически низкий'}

{llm_section}

ОБЩАЯ ИНФОРМАЦИЯ
----------------
• Код ответа: {safe_get(result, ['basic_info', 'status_code'])}
• Время загрузки: {safe_get(result, ['basic_info', 'response_time'])} сек
• Размер страницы: {safe_get(result, ['performance', 'page_size'], 0) // 1024} KB
• Использует HTTPS: {'Да' if safe_get(result, ['basic_info', 'is_https'], False) else 'Нет'}
• Редиректов: {safe_get(result, ['basic_info', 'redirects'], 0)}
• Домен: {safe_get(result, ['basic_info', 'domain'], 'Не доступно')}

КРИТИЧЕСКИЕ ПРОБЛЕМЫ ({len(safe_get(result, ['critical_issues'], []))})
---------------------
{chr(10).join(['• ' + issue for issue in safe_get(result, ['critical_issues'], [])]) if safe_get(result, ['critical_issues'], []) else '• Отсутствуют'}

ПРЕДУПРЕЖДЕНИЯ ({len(safe_get(result, ['warnings'], []))})
---------------
{chr(10).join(['• ' + warning for warning in safe_get(result, ['warnings'], [])]) if safe_get(result, ['warnings'], []) else '• Отсутствуют'}

РЕКОМЕНДАЦИИ ({len(safe_get(result, ['recommendations'], []))})
-------------
{chr(10).join(['• ' + str(rec) for rec in safe_get(result, ['recommendations'], [])]) if safe_get(result, ['recommendations'], []) else '• Все рекомендации выполнены'}

ДЕТАЛЬНЫЙ АНАЛИЗ
===============

1. МЕТА-ДАННЫЕ
-------------
Title:
  • Наличие: {'Да' if safe_get(result, ['metadata', 'title', 'value']) else 'Нет'}
  • Длина: {safe_get(result, ['metadata', 'title', 'length'], 0)} символов
  • Оптимальность: {'Оптимально (50-60 символов)' if safe_get(result, ['metadata', 'title', 'optimal'], False) else 'Не оптимально'}
  • Содержимое: {safe_get(result, ['metadata', 'title', 'value'], 'Отсутствует') or 'Отсутствует'}

Description:
  • Наличие: {'Да' if safe_get(result, ['metadata', 'description', 'value']) else 'Нет'}
  • Длина: {safe_get(result, ['metadata', 'description', 'length'], 0)} символов
  • Оптимальность: {'Оптимально (120-160 символов)' if safe_get(result, ['metadata', 'description', 'optimal'], False) else 'Не оптимально'}
  • Содержимое: {safe_get(result, ['metadata', 'description', 'value'], 'Отсутствует') or 'Отсутствует'}

Open Graph:
  • Наличие: {'Да' if safe_get(result, ['metadata', 'open_graph', 'exists'], False) else 'Нет'}
  • Количество тегов: {safe_get(result, ['metadata', 'open_graph', 'count'], 0)}
  • Основные теги: {safe_get(result, ['metadata', 'open_graph', 'essential_count'], 0)}/4

Twitter Cards:
  • Наличие: {'Да' if safe_get(result, ['metadata', 'twitter_cards', 'exists'], False) else 'Нет'}
  • Количество тегов: {safe_get(result, ['metadata', 'twitter_cards', 'count'], 0)}

Технические мета-теги:
  • Canonical: {'Присутствует' if safe_get(result, ['metadata', 'canonical', 'exists'], False) else 'Отсутствует'}
  • Robots: {'Присутствует' if safe_get(result, ['metadata', 'robots', 'exists'], False) else 'Отсутствует'}
  • Viewport: {'Присутствует' if safe_get(result, ['metadata', 'viewport', 'exists'], False) else 'Отсутствует'}
  • Charset: {'Присутствует' if safe_get(result, ['metadata', 'charset', 'exists'], False) else 'Отсутствует'}

2. СЕМАНТИЧЕСКАЯ РАЗМЕТКА
------------------------
Schema.org (JSON-LD):
  • Наличие: {'Да' if safe_get(result, ['semantic_markup', 'schema_org', 'exists'], False) else 'Нет'}
  • Количество скриптов: {safe_get(result, ['semantic_markup', 'schema_org', 'scripts'], 0)}
  • Типы разметки: {', '.join(safe_get(result, ['semantic_markup', 'schema_org', 'types'], [])) if safe_get(result, ['semantic_markup', 'schema_org', 'types'], []) else 'Не определены'}

Микроразметка (Microdata):
  • Наличие: {'Да' if safe_get(result, ['semantic_markup', 'microdata', 'exists'], False) else 'Нет'}
  • Количество элементов: {safe_get(result, ['semantic_markup', 'microdata', 'elements'], 0)}

RDFa:
  • Наличие: {'Да' if safe_get(result, ['semantic_markup', 'rdfa', 'exists'], False) else 'Нет'}
  • Количество элементов: {safe_get(result, ['semantic_markup', 'rdfa', 'elements'], 0)}

Заголовки:
  • H1: {safe_get(result, ['semantic_markup', 'headings', 'h1'], 0)}
  • H2: {safe_get(result, ['semantic_markup', 'headings', 'h2'], 0)}
  • H3: {safe_get(result, ['semantic_markup', 'headings', 'h3'], 0)}
  • H4: {safe_get(result, ['semantic_markup', 'headings', 'h4'], 0)}
  • H5: {safe_get(result, ['semantic_markup', 'headings', 'h5'], 0)}
  • H6: {safe_get(result, ['semantic_markup', 'headings', 'h6'], 0)}

Иерархия заголовков:
  • Один H1: {'Да' if safe_get(result, ['semantic_markup', 'heading_hierarchy', 'has_single_h1'], False) else 'Нет'}
  • Корректная иерархия: {'Да' if safe_get(result, ['semantic_markup', 'heading_hierarchy', 'hierarchy_correct'], False) else 'Нет'}

3. СТРУКТУРА КОНТЕНТА
--------------------
Объем контента:
  • Количество слов: {safe_get(result, ['content_structure', 'word_count'], 0)}
  • Уникальные слова: {safe_get(result, ['content_structure', 'keyword_analysis', 'unique_words'], 0)}
  • Соотношение текст/HTML: {safe_get(result, ['content_structure', 'text_ratio'], 0)}%
  • Плотность ключевых слов: {safe_get(result, ['content_structure', 'keyword_analysis', 'keyword_density'], 0)}%

Элементы структуры:
  • Маркированные списки (ul): {safe_get(result, ['content_structure', 'lists', 'ul'], 0)}
  • Нумерованные списки (ol): {safe_get(result, ['content_structure', 'lists', 'ol'], 0)}
  • Всего списков: {safe_get(result, ['content_structure', 'lists', 'total'], 0)}
  • Таблицы: {safe_get(result, ['content_structure', 'tables'], 0)}

Изображения:
  • Всего изображений: {safe_get(result, ['content_structure', 'images', 'total'], 0)}
  • С alt-текстом: {safe_get(result, ['content_structure', 'images', 'with_alt'], 0)}
  • Процент с alt-текстом: {safe_get(result, ['content_structure', 'images', 'alt_percentage'], 0)}%

Мультимедиа:
  • Видео: {safe_get(result, ['content_structure', 'multimedia', 'videos'], 0)}
  • Iframes: {safe_get(result, ['content_structure', 'multimedia', 'iframes'], 0)}
  • Аудио: {safe_get(result, ['content_structure', 'multimedia', 'audio'], 0)}

Интерактивные элементы:
  • Формы: {safe_get(result, ['content_structure', 'interactive', 'forms'], 0)}
  • Кнопки: {safe_get(result, ['content_structure', 'interactive', 'buttons'], 0)}
  • Поля ввода: {safe_get(result, ['content_structure', 'interactive', 'inputs'], 0)}

Анализ читаемости:
  • Оценка читаемости: {safe_get(result, ['content_structure', 'readability', 'score'], 0)}/100
  • Уровень: {safe_get(result, ['content_structure', 'readability', 'level'], 'Не доступно')}
  • Средняя длина предложения: {safe_get(result, ['content_structure', 'readability', 'avg_sentence_length'], 0)} слов
  • Средняя длина слова: {safe_get(result, ['content_structure', 'readability', 'avg_word_length'], 0)} символов
  • Абзацы: {safe_get(result, ['content_structure', 'readability', 'paragraphs'], 0)}

Ключевые слова (топ-15):
{chr(10).join(['  • ' + word[0] + ': ' + str(word[1]) + ' раз' for word in safe_get(result, ['content_structure', 'keyword_analysis', 'top_words'], [])]) if safe_get(result, ['content_structure', 'keyword_analysis', 'top_words'], []) else '  • Анализ не выполнен'}

4. ТЕХНИЧЕСКИЕ АСПЕКТЫ
---------------------
Ссылки:
  • Всего ссылок: {safe_get(result, ['technical_seo', 'links', 'total'], 0)}
  • Внутренние ссылки: {safe_get(result, ['technical_seo', 'links', 'internal'], 0)}
  • Внешние ссылки: {safe_get(result, ['technical_seo', 'links', 'external'], 0)}
  • Ссылки с анкором: {safe_get(result, ['technical_seo', 'links', 'with_anchor'], 0)}

Улучшенный анализ ссылок:
  • Всего ссылок: {safe_get(result, ['technical_seo', 'enhanced_links', 'total'], 0)}
  • Внутренние: {safe_get(result, ['technical_seo', 'enhanced_links', 'internal'], 0)}
  • Внешние: {safe_get(result, ['technical_seo', 'enhanced_links', 'external'], 0)}
  • Nofollow: {safe_get(result, ['technical_seo', 'enhanced_links', 'nofollow'], 0)}
  • Dofollow: {safe_get(result, ['technical_seo', 'enhanced_links', 'dofollow'], 0)}
  • С анкор-текстом: {safe_get(result, ['technical_seo', 'enhanced_links', 'with_anchor'], 0)}
  • Без анкор-текста: {safe_get(result, ['technical_seo', 'enhanced_links', 'empty_anchor'], 0)}

Анализ изображений:
  • Всего изображений: {safe_get(result, ['technical_seo', 'images_analysis', 'total'], 0)}
  • С alt-текстом: {safe_get(result, ['technical_seo', 'images_analysis', 'with_alt'], 0)}
  • Lazy loading: {safe_get(result, ['technical_seo', 'images_analysis', 'lazy_loaded'], 0)}
  • Responsive: {safe_get(result, ['technical_seo', 'images_analysis', 'responsive'], 0)}

Структура URL:
  • Глубина: {safe_get(result, ['technical_seo', 'url_structure', 'depth'], 0)} уровней
  • Завершающий слэш: {'Да' if safe_get(result, ['technical_seo', 'url_structure', 'has_trailing_slash'], False) else 'Нет'}
  • Параметры: {'Да' if safe_get(result, ['technical_seo', 'url_structure', 'has_parameters'], False) else 'Нет'}

Важные технические теги:
  • Robots.txt: {'Найден' if safe_get(result, ['technical_seo', 'important_tags', 'robots_txt'], False) else 'Не найден'}
  • Sitemap.xml: {'Найден' if safe_get(result, ['technical_seo', 'important_tags', 'sitemap'], False) else 'Не найден'}
  • Favicon: {'Найден' if safe_get(result, ['technical_seo', 'important_tags', 'favicon'], False) else 'Не найден'}
  • Web App Manifest: {'Найден' if safe_get(result, ['technical_seo', 'important_tags', 'manifest'], False) else 'Не найден'}
  • AMP: {'Найден' if safe_get(result, ['technical_seo', 'important_tags', 'amp_html'], False) else 'Не найден'}

5. ПРОИЗВОДИТЕЛЬНОСТЬ
--------------------
• Оценка производительности: {safe_get(result, ['performance', 'score'], 0)}/100
• Уровень: {safe_get(result, ['performance', 'level'], 'Не доступно')}
• Время ответа сервера: {safe_get(result, ['performance', 'response_time'], 0)} сек
• Размер HTML: {safe_get(result, ['performance', 'html_size'], 0)} байт
• Размер страницы: {safe_get(result, ['performance', 'page_size'], 0)} байт
• Сложность HTML: {safe_get(result, ['performance', 'html_complexity'], 0)}
• Элементов DOM: {safe_get(result, ['performance', 'dom_elements'], 0)}
• Глубина DOM: {safe_get(result, ['performance', 'dom_depth'], 0)}
• Количество изображений: {safe_get(result, ['performance', 'image_count'], 0)}
• Количество скриптов: {safe_get(result, ['performance', 'script_count'], 0)}
• Количество стилей: {safe_get(result, ['performance', 'stylesheet_count'], 0)}
• Внешние скрипты: {safe_get(result, ['performance', 'external_scripts'], 0)}
• Встроенные скрипты: {safe_get(result, ['performance', 'inline_scripts'], 0)}
• Запросы ресурсов: {safe_get(result, ['performance', 'resource_requests'], 0)}

6. БЕЗОПАСНОСТЬ
---------------
HTTPS:
  • Включен: {'Да' if safe_get(result, ['security', 'https', 'enabled'], False) else 'Нет'}
  • Смешанный контент: {'Обнаружен' if safe_get(result, ['security', 'https', 'mixed_content'], False) else 'Не обнаружен'}

Заголовки безопасности:
  • HSTS: {'Включен' if safe_get(result, ['security', 'headers', 'hsts'], False) else 'Отсутствует'}
  • X-Frame-Options: {'Включен' if safe_get(result, ['security', 'headers', 'x_frame_options'], False) else 'Отсутствует'}
  • X-Content-Type-Options: {'Включен' if safe_get(result, ['security', 'headers', 'x_content_type_options'], False) else 'Отсутствует'}
  • X-XSS-Protection: {'Включен' if safe_get(result, ['security', 'headers', 'x_xss_protection'], False) else 'Отсутствует'}
  • Content-Security-Policy: {'Включен' if safe_get(result, ['security', 'headers', 'content_security_policy'], False) else 'Отсутствует'}
  • Referrer-Policy: {'Включен' if safe_get(result, ['security', 'headers', 'referrer_policy'], False) else 'Отсутствует'}

7. ДОСТУПНОСТЬ
--------------
ARIA атрибуты:
  • Метки: {safe_get(result, ['accessibility', 'aria', 'labels'], 0)}
  • Роли: {safe_get(result, ['accessibility', 'aria', 'roles'], 0)}
  • Описания: {safe_get(result, ['accessibility', 'aria', 'describedby'], 0)}

Семантические теги:
  • Header: {safe_get(result, ['accessibility', 'semantic_html', 'header'], 0)}
  • Footer: {safe_get(result, ['accessibility', 'semantic_html', 'footer'], 0)}
  • Nav: {safe_get(result, ['accessibility', 'semantic_html', 'nav'], 0)}
  • Main: {safe_get(result, ['accessibility', 'semantic_html', 'main'], 0)}
  • Article: {safe_get(result, ['accessibility', 'semantic_html', 'article'], 0)}
  • Section: {safe_get(result, ['accessibility', 'semantic_html', 'section'], 0)}
  • Aside: {safe_get(result, ['accessibility', 'semantic_html', 'aside'], 0)}

Формы:
  • Всего форм: {safe_get(result, ['accessibility', 'forms', 'total'], 0)}
  • С метками: {safe_get(result, ['accessibility', 'forms', 'with_labels'], 0)}
  • С плейсхолдерами: {safe_get(result, ['accessibility', 'forms', 'with_placeholders'], 0)}

{charts_html}

ПРИОРИТЕТНЫЕ РЕКОМЕНДАЦИИ
=========================

КРИТИЧЕСКИЕ ИСПРАВЛЕНИЯ:
{chr(10).join(['• ' + issue for issue in safe_get(result, ['critical_issues'], [])]) if safe_get(result, ['critical_issues'], []) else '• Отсутствуют'}

ВАЖНЫЕ УЛУЧШЕНИЯ:
{chr(10).join(['• ' + warning for warning in safe_get(result, ['warnings'], [])]) if safe_get(result, ['warnings'], []) else '• Отсутствуют'}

ДОПОЛНИТЕЛЬНЫЕ РЕКОМЕНДАЦИИ:
{generate_additional_recommendations(result)}

GEO-РЕКОМЕНДАЦИИ ОТ AI-МОДЕЛЕЙ
==============================
{chr(10).join(['• ' + rec for rec in safe_get(result, ['llm_analysis', 'geo_recommendations'], [])]) if safe_get(result, ['llm_analysis', 'geo_recommendations'], []) else '• Специфические GEO-рекомендации отсутствуют'}

---
Отчет сгенерирован GEO Analyzer Pro с использованием AI-моделей
{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
    """
    
    return report

def generate_llm_analysis_section(result):
    """Генерация раздела LLM-анализа с полными текстами ответов"""
    llm_analysis = safe_get(result, ['llm_analysis'], {})
    
    if not llm_analysis or 'error' in llm_analysis:
        return """
AI-АНАЛИЗ ДЛЯ ГЕНЕРАТИВНОГО ПОИСКА
----------------------------------
LLM-анализ временно недоступен. Используются стандартные метрики оптимизации.
        """
    
    models_used = safe_get(llm_analysis, ['models_used'], [])
    geo_score = safe_get(llm_analysis, ['overall_geo_score'], 0)
    citation_potential = safe_get(llm_analysis, ['citation_potential'], 0)
    analysis_summary = safe_get(llm_analysis, ['analysis_summary'], '')
    
    section = f"""
AI-АНАЛИЗ ДЛЯ ГЕНЕРАТИВНОГО ПОИСКА
----------------------------------
Использованные AI-модели: {', '.join(models_used) if models_used else 'Не доступно'}
Общая оценка GEO: {geo_score}/100
Потенциал цитирования: {citation_potential}/100

СВОДКА АНАЛИЗА:
{analysis_summary}

ИНСАЙТЫ ОТ AI-МОДЕЛЕЙ:
{chr(10).join(['• ' + insight for insight in safe_get(llm_analysis, ['llm_insights'], [])]) if safe_get(llm_analysis, ['llm_insights'], []) else '• Инсайты не сгенерированы'}

ПОЛНЫЕ ОТВЕТЫ ОТ AI-МОДЕЛЕЙ:
"""
    
    # Добавляем полные тексты ответов от каждой модели
    detailed_analysis = safe_get(llm_analysis, ['detailed_analysis'], {})
    raw_responses = safe_get(llm_analysis, ['raw_responses'], {})
    
    for model in models_used:
        model_display_name = {
            'bert_nebulon': 'BERT-NEBULON ALPHA',
            'grok': 'GROK 4.1 FAST',
            'deepseek': 'DEEPSEEK R1T2 CHIMERA'
        }.get(model, model.upper())
        
        section += f"\n{'='*60}\n"
        section += f"{model_display_name} - ПОЛНЫЙ ОТВЕТ\n"
        section += f"{'='*60}\n\n"
        
        # Сначала показываем подробный анализ
        if model in detailed_analysis and detailed_analysis[model]:
            section += f"ДЕТАЛЬНЫЙ АНАЛИЗ:\n{detailed_analysis[model]}\n\n"
        
        # Затем показываем полный сырой ответ
        if model in raw_responses and raw_responses[model]:
            section += f"ПОЛНЫЙ ТЕКСТ ОТВЕТА:\n{raw_responses[model]}\n"
        else:
            section += "Полный текст ответа модели не доступен.\n"
        
        section += "\n" + "-"*60 + "\n"
    
    return section
  
def generate_charts_html(result):
    """Генерация HTML с графиками для отчета"""
    try:
        charts = []
        
        # Создаем график распределения оценок по категориям
        categories = ['Мета-данные', 'Семантика', 'Контент', 'Техника', 'Производительность', 'Безопасность', 'Доступность']
        
        # Безопасное получение значений с проверкой на существование ключей
        overall_score = safe_get(result, ['score'], 0)
        performance_score = safe_get(result, ['performance', 'score'], 0)
        security_score = 5 if safe_get(result, ['security', 'https', 'enabled'], False) else 0
        semantic_html = safe_get(result, ['accessibility', 'semantic_html'], {})
        accessibility_score = 5 if semantic_html and any(semantic_html.values()) else 0
        
        scores = [
            min(20, overall_score), 
            min(20, overall_score), 
            min(20, overall_score),
            min(20, overall_score),
            performance_score,
            security_score,
            accessibility_score
        ]
        
        fig = go.Figure(data=[
            go.Bar(name='Текущие оценки', x=categories, y=scores,
                   marker_color=['#2962FF', '#00C853', '#FFD600', '#FF9100', '#D50000', '#6200EA', '#2196F3'])
        ])
        
        fig.update_layout(
            title='Оценки по категориям',
            yaxis_title='Оценка',
            yaxis_range=[0, 20]
        )
        
        charts.append(fig.to_html(full_html=False))
        
        return "\n".join([f"<div style='page-break-before: always;'>{chart}</div>" for chart in charts])
    except Exception as e:
        return f"<!-- Ошибка генерации графиков: {e} -->"
      
def generate_additional_recommendations(result):
    """Генерация дополнительных рекомендаций с защитой от отсутствующих полей"""
    recommendations = []
    
    # Безопасное получение значений
    word_count = safe_get(result, ['content_structure', 'word_count'], 0)
    lists_total = safe_get(result, ['content_structure', 'lists', 'total'], 0)
    tables_count = safe_get(result, ['content_structure', 'tables'], 0)
    schema_org_exists = safe_get(result, ['semantic_markup', 'schema_org', 'exists'], False)
    microdata_elements = safe_get(result, ['semantic_markup', 'microdata', 'elements'], 0)
    response_time = safe_get(result, ['performance', 'response_time'], 0)
    page_size = safe_get(result, ['performance', 'page_size'], 0)
    https_enabled = safe_get(result, ['security', 'https', 'enabled'], False)
    mixed_content = safe_get(result, ['security', 'https', 'mixed_content'], False)
    aria_labels = safe_get(result, ['accessibility', 'aria', 'labels'], 0)
    semantic_html_values = safe_get(result, ['accessibility', 'semantic_html'], {})
    
    # Рекомендации по контенту
    if word_count < OPTIMAL_VALUES['word_count_good']:
        recommendations.append(f"Увеличить объем контента до {OPTIMAL_VALUES['word_count_good']}+ слов для лучшего покрытия темы")
    
    if lists_total < 2:
        recommendations.append("Добавить больше списков для улучшения структуры контента")
    
    if tables_count == 0:
        recommendations.append("Рассмотреть добавление таблиц для структурирования данных")
    
    # Рекомендации по семантике
    if not schema_org_exists:
        recommendations.append("Добавить Schema.org разметку для улучшения понимания контента поисковыми системами")
    
    if microdata_elements == 0:
        recommendations.append("Внедрить микроразметку для ключевых элементов страницы")
    
    # Рекомендации по производительности
    if response_time > 2:
        recommendations.append("Оптимизировать время загрузки страницы (цель < 2 секунд)")
    
    if page_size > OPTIMAL_VALUES['page_size_max']:
        recommendations.append("Уменьшить размер страницы за счет оптимизации изображений и кода")
    
    # Рекомендации по безопасности
    if not https_enabled:
        recommendations.append("Перевести сайт на HTTPS для обеспечения безопасности пользователей")
    
    if mixed_content:
        recommendations.append("Исправить смешанный контент, заменив HTTP ресурсы на HTTPS")
    
    # Рекомендации по доступности
    if aria_labels == 0:
        recommendations.append("Добавить ARIA-метки для улучшения доступности сайта")
    
    if semantic_html_values and sum(semantic_html_values.values()) == 0:
        recommendations.append("Использовать семантические HTML5 теги для улучшения структуры страницы")
    
    """Генерация дополнительных рекомендаций с примерами"""
    additional_recs = []
    
    return "\n".join(additional_recs) if additional_recs else "• Все ключевые аспекты оптимизированы!"
  
def create_download_link(report, filename="geo_analysis_report.txt"):
    """Создание ссылки для скачивания отчета"""
    b64 = base64.b64encode(report.encode()).decode()
    href = f'<a href="data:file/txt;base64,{b64}" download="{filename}" class="quick-action-btn">📥 Скачать полный отчет (.txt)</a>'
    return href

def generate_docx_report(result):
    """Генерация отчета в формате DOCX"""
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document()
        
        # Заголовок документа
        title = doc.add_heading('GEO ANALYZER PRO - ПОЛНЫЙ ОТЧЕТ АНАЛИЗА САЙТА', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Основная информация
        doc.add_heading('Общая информация', level=1)
        
        basic_info = [
            ("Дата анализа", result['basic_info']['analysis_date']),
            ("Анализируемый URL", result['basic_info']['url']),
            ("Финальный URL", result['basic_info'].get('final_url', result['basic_info']['url'])),
            ("Общая оценка", f"{result['score']}/100"),
            ("Статус", 'Отлично' if result['score'] >= 80 else 'Хорошо' if result['score'] >= 60 else 'Требует улучшений' if result['score'] >= 40 else 'Критически низкий'),
            ("Код ответа", str(result['basic_info']['status_code'])),
            ("Время загрузки", f"{result['basic_info']['response_time']} сек"),
            ("Размер страницы", f"{result['performance']['page_size'] // 1024} KB"),
            ("Использует HTTPS", 'Да' if result['basic_info']['is_https'] else 'Нет'),
            ("Редиректов", str(result['basic_info']['redirects']))
        ]
        
        for label, value in basic_info:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(str(value))
        
        # Критические проблемы
        doc.add_heading('Критические проблемы', level=1)
        if result['critical_issues']:
            for issue in result['critical_issues']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(issue).bold = True
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Красный цвет
        else:
            doc.add_paragraph('• Отсутствуют')
        
        # Предупреждения
        doc.add_heading('Предупреждения', level=1)
        if result['warnings']:
            for warning in result['warnings']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(warning).bold = True
                p.runs[0].font.color.rgb = RGBColor(255, 165, 0)  # Оранжевый цвет
        else:
            doc.add_paragraph('• Отсутствуют')
        
        # Рекомендации
        doc.add_heading('Рекомендации', level=1)
        if result['recommendations']:
            for rec in result['recommendations']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(rec)
        else:
            doc.add_paragraph('• Все рекомендации выполнены')
        
        # Детальный анализ
        doc.add_heading('Детальный анализ', level=1)
        
        # Мета-данные
        doc.add_heading('1. Мета-данные', level=2)
        metadata = result['metadata']
        
        doc.add_heading('Title:', level=3)
        p = doc.add_paragraph()
        p.add_run('• Наличие: ').bold = True
        p.add_run('Да' if metadata['title']['value'] else 'Нет')
        
        p = doc.add_paragraph()
        p.add_run('• Длина: ').bold = True
        p.add_run(f"{metadata['title']['length']} символов")
        
        p = doc.add_paragraph()
        p.add_run('• Оптимальность: ').bold = True
        p.add_run('Оптимально (50-60 символов)' if metadata['title']['optimal'] else 'Не оптимально')
        
        p = doc.add_paragraph()
        p.add_run('• Содержимое: ').bold = True
        p.add_run(metadata['title']['value'] or 'Отсутствует')
        
        # Семантическая разметка
        doc.add_heading('2. Семантическая разметка', level=2)
        semantic = result['semantic_markup']
        
        doc.add_heading('Schema.org (JSON-LD):', level=3)
        p = doc.add_paragraph()
        p.add_run('• Наличие: ').bold = True
        p.add_run('Да' if semantic['schema_org']['exists'] else 'Нет')
        
        p = doc.add_paragraph()
        p.add_run('• Количество скриптов: ').bold = True
        p.add_run(str(semantic['schema_org']['scripts']))
        
        # Структура контента
        doc.add_heading('3. Структура контента', level=2)
        content = result['content_structure']
        
        doc.add_heading('Объем контента:', level=3)
        p = doc.add_paragraph()
        p.add_run('• Количество слов: ').bold = True
        p.add_run(str(content['word_count']))
        
        p = doc.add_paragraph()
        p.add_run('• Соотношение текст/HTML: ').bold = True
        p.add_run(f"{content['text_ratio']}%")
        
        # Технические аспекты
        doc.add_heading('4. Технические аспекты', level=2)
        technical = result['technical_seo']
        
        doc.add_heading('Ссылки:', level=3)
        p = doc.add_paragraph()
        p.add_run('• Всего ссылок: ').bold = True
        p.add_run(str(technical['links']['total']))
        
        p = doc.add_paragraph()
        p.add_run('• Внутренние ссылки: ').bold = True
        p.add_run(str(technical['links']['internal']))
        
        p = doc.add_paragraph()
        p.add_run('• Внешние ссылки: ').bold = True
        p.add_run(str(technical['links']['external']))
        
        # Производительность
        doc.add_heading('5. Производительность', level=2)
        performance = result['performance']
        
        p = doc.add_paragraph()
        p.add_run('• Оценка производительности: ').bold = True
        p.add_run(f"{performance['score']}/100")
        
        p = doc.add_paragraph()
        p.add_run('• Время ответа сервера: ').bold = True
        p.add_run(f"{performance['response_time']} сек")
        
        # Безопасность
        doc.add_heading('6. Безопасность', level=2)
        security = result['security']
        
        p = doc.add_paragraph()
        p.add_run('• HTTPS: ').bold = True
        p.add_run('Включен' if security['https']['enabled'] else 'Отключен')
        
        # Доступность
        doc.add_heading('7. Доступность', level=2)
        accessibility = result['accessibility']
        
        p = doc.add_paragraph()
        p.add_run('• ARIA метки: ').bold = True
        p.add_run(str(accessibility['aria']['labels']))
        
        # Рекомендации
        doc.add_heading('Приоритетные рекомендации', level=1)
        
        doc.add_heading('Критические исправления:', level=2)
        if result['critical_issues']:
            for issue in result['critical_issues']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(issue).bold = True
        else:
            doc.add_paragraph('• Отсутствуют')
        
        doc.add_heading('Важные улучшения:', level=2)
        if result['warnings']:
            for warning in result['warnings']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(warning).bold = True
        else:
            doc.add_paragraph('• Отсутствуют')
        
        doc.add_heading('Дополнительные рекомендации:', level=2)
        additional_recs = generate_additional_recommendations(result).split('\n')
        for rec in additional_recs:
            if rec.strip() and not rec.startswith("• Все ключевые"):
                doc.add_paragraph(rec, style='List Bullet')
        
        # Футер
        doc.add_paragraph('---')
        p = doc.add_paragraph('Отчет сгенерирован GEO Analyzer Pro')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return doc
        
    except Exception as e:
        return None

def create_docx_download_link(doc, filename="geo_analysis_report.docx"):
    """Создание ссылки для скачивания DOCX отчета"""
    if not DOCX_AVAILABLE:
        return '<p style="color: red;">⚠️ Библиотека python-docx не установлена</p>'
    
    try:
        # Сохраняем документ в байтовый поток
        import io
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        # Кодируем в base64
        b64 = base64.b64encode(doc_bytes.getvalue()).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" class="quick-action-btn">📄 Скачать полный отчет (.docx)</a>'
        return href
        
    except Exception as e:
        return f'<p style="color: red;">❌ Ошибка создания DOCX: {str(e)}</p>'

def generate_comparison_report(current_result, previous_result):
    """Генерация сравнительного отчета"""
    if not previous_result:
        return "Нет данных для сравнения"
    
    comparison = f"""
СРАВНИТЕЛЬНЫЙ ОТЧЕТ GEO ANALYZER PRO
====================================
Период сравнения: {previous_result['basic_info']['analysis_date']} -> {current_result['basic_info']['analysis_date']}

ОБЩАЯ ОЦЕНКА
------------
Текущая оценка: {current_result['score']}/100
Предыдущая оценка: {previous_result['score']}/100
Изменение: {current_result['score'] - previous_result['score']} баллов

КЛЮЧЕВЫЕ ИЗМЕНЕНИЯ
------------------

Мета-данные:
  • Title: {previous_result['metadata']['title']['length']} -> {current_result['metadata']['title']['length']} символов
  • Description: {previous_result['metadata']['description']['length']} -> {current_result['metadata']['description']['length']} символов

Контент:
  • Количество слов: {previous_result['content_structure']['word_count']} -> {current_result['content_structure']['word_count']}
  • Alt-тексты: {previous_result['content_structure']['images']['alt_percentage']}% -> {current_result['content_structure']['images']['alt_percentage']}%

Производительность:
  • Время загрузки: {previous_result['performance']['response_time']}с -> {current_result['performance']['response_time']}с
  • Размер страницы: {previous_result['performance']['page_size'] // 1024} KB -> {current_result['performance']['page_size'] // 1024} KB

ВЫВОДЫ
------
"""
    
    score_diff = current_result['score'] - previous_result['score']
    if score_diff > 0:
        comparison += f"✅ Улучшение на {score_diff} баллов. Продолжайте в том же духе!"
    elif score_diff < 0:
        comparison += f"⚠️ Снижение на {abs(score_diff)} баллов. Обратите внимание на ухудшившиеся показатели."
    else:
        comparison += "➡️ Без изменений. Рассмотрите возможность оптимизации дополнительных параметров."
    
    return comparison